"""Tests for file parsers (VTT, DOCX, DOC, PDF, Markdown)."""

import io
from unittest.mock import MagicMock, patch

import pytest

from backend.parsers import (
    ALLOWED_EXTENSIONS,
    MAX_FILE_SIZE,
    parse_docx,
    parse_doc,
    parse_file,
    parse_markdown,
    parse_pdf,
    parse_vtt,
)


# ── VTT tests ───────────────────────────────────────────────────────────


class TestParseVtt:
    def _make_vtt(self, cues: str) -> bytes:
        return f"WEBVTT\n\n{cues}".encode("utf-8")

    def test_teams_vtt_with_speaker_tags(self):
        vtt = self._make_vtt(
            "00:00:01.000 --> 00:00:05.000\n"
            "<v Alice>Hello everyone</v>\n\n"
            "00:00:06.000 --> 00:00:10.000\n"
            "<v Bob>Hi Alice</v>\n"
        )
        result = parse_vtt(vtt)

        assert result.format == "vtt"
        assert "[00:00:01] Alice: Hello everyone" in result.text
        assert "[00:00:06] Bob: Hi Alice" in result.text

    def test_merge_consecutive_same_speaker(self):
        vtt = self._make_vtt(
            "00:00:01.000 --> 00:00:03.000\n"
            "<v Alice>First part</v>\n\n"
            "00:00:03.000 --> 00:00:06.000\n"
            "<v Alice>second part</v>\n"
        )
        result = parse_vtt(vtt)

        # Should merge into one line
        assert result.text.count("Alice:") == 1
        assert "First part second part" in result.text

    def test_no_speaker_fallback(self):
        vtt = self._make_vtt(
            "00:00:01.000 --> 00:00:05.000\n"
            "Plain caption text\n"
        )
        result = parse_vtt(vtt)

        assert "[00:00:01] Plain caption text" in result.text

    def test_empty_cues_skipped(self):
        vtt = self._make_vtt(
            "00:00:01.000 --> 00:00:05.000\n"
            "<v Alice>Hello</v>\n\n"
            "00:00:06.000 --> 00:00:10.000\n"
            "\n"
        )
        result = parse_vtt(vtt)

        assert "Hello" in result.text
        lines = [l for l in result.text.splitlines() if l.strip()]
        assert len(lines) == 1

    def test_utf8_bom_handled(self):
        vtt = b"\xef\xbb\xbfWEBVTT\n\n00:00:01.000 --> 00:00:05.000\n<v Alice>Hello</v>\n"
        result = parse_vtt(vtt)

        assert "Alice: Hello" in result.text


# ── DOCX tests ──────────────────────────────────────────────────────────


class TestParseDocx:
    def _make_docx(self, paragraphs: list[str]) -> bytes:
        from docx import Document

        doc = Document()
        for p in paragraphs:
            doc.add_paragraph(p)
        buf = io.BytesIO()
        doc.save(buf)
        return buf.getvalue()

    def test_teams_docx_speaker_pattern(self):
        content = self._make_docx([
            "Alice Johnson  0:15",
            "Hello everyone, let's start.",
            "Bob Smith  1:30",
            "I have an update on the API.",
        ])
        result = parse_docx(content)

        assert result.format == "docx"
        assert "[00:00:15] Alice Johnson: Hello everyone" in result.text
        assert "[00:01:30] Bob Smith: I have an update" in result.text

    def test_generic_docx_fallback(self):
        content = self._make_docx([
            "Meeting Notes",
            "We discussed the project timeline.",
            "Next steps were defined.",
        ])
        result = parse_docx(content)

        assert "Meeting Notes" in result.text
        assert "We discussed" in result.text
        assert result.detected_title == "Meeting Notes"

    def test_hhmmss_timestamp(self):
        content = self._make_docx([
            "Speaker  1:02:30",
            "Long meeting comment.",
        ])
        result = parse_docx(content)

        assert "[01:02:30] Speaker: Long meeting comment." in result.text

    def test_empty_docx(self):
        content = self._make_docx([])
        result = parse_docx(content)

        assert result.text == ""


# ── DOC tests ───────────────────────────────────────────────────────────


class TestParseDoc:
    def test_antiword_success(self):
        mock_result = MagicMock()
        mock_result.returncode = 0
        mock_result.stdout = "Speaker: Hello from doc file"
        mock_result.stderr = ""

        with patch("backend.parsers.subprocess.run", return_value=mock_result):
            result = parse_doc(b"fake doc content")

        assert result.format == "doc"
        assert "Hello from doc file" in result.text

    def test_antiword_failure_raises(self):
        mock_result = MagicMock()
        mock_result.returncode = 1
        mock_result.stderr = "not a word document"

        with patch("backend.parsers.subprocess.run", return_value=mock_result):
            with pytest.raises(RuntimeError, match="antiword failed"):
                parse_doc(b"not a real doc")

    def test_antiword_not_found(self):
        with patch(
            "backend.parsers.subprocess.run",
            side_effect=FileNotFoundError("antiword not found"),
        ):
            with pytest.raises(FileNotFoundError):
                parse_doc(b"fake content")

    def test_temp_file_cleaned_up(self):
        mock_result = MagicMock()
        mock_result.returncode = 0
        mock_result.stdout = "text"
        mock_result.stderr = ""

        with patch("backend.parsers.subprocess.run", return_value=mock_result) as mock_run:
            parse_doc(b"content")
            # Verify the temp file path was passed and would be cleaned up
            call_args = mock_run.call_args
            tmp_path = call_args[0][0][1]
            import os
            assert not os.path.exists(tmp_path)


# ── PDF tests ───────────────────────────────────────────────────────────


class TestParsePdf:
    def _make_pdf(self, text: str) -> bytes:
        from pypdf import PdfWriter

        writer = PdfWriter()
        # pypdf doesn't easily create text pages from scratch,
        # so we mock instead for reliable testing
        return b""  # placeholder

    def test_text_pdf(self):
        mock_reader = MagicMock()
        mock_page = MagicMock()
        mock_page.extract_text.return_value = "Page 1 content"
        mock_reader.pages = [mock_page]

        with patch("backend.parsers.PdfReader", return_value=mock_reader):
            result = parse_pdf(b"fake pdf")

        assert result.format == "pdf"
        assert "Page 1 content" in result.text

    def test_multi_page_pdf(self):
        mock_reader = MagicMock()
        pages = []
        for i in range(3):
            p = MagicMock()
            p.extract_text.return_value = f"Page {i + 1}"
            pages.append(p)
        mock_reader.pages = pages

        with patch("backend.parsers.PdfReader", return_value=mock_reader):
            result = parse_pdf(b"fake pdf")

        assert "Page 1" in result.text
        assert "Page 2" in result.text
        assert "Page 3" in result.text

    def test_empty_pdf(self):
        mock_reader = MagicMock()
        mock_reader.pages = []

        with patch("backend.parsers.PdfReader", return_value=mock_reader):
            result = parse_pdf(b"fake pdf")

        assert result.text == ""

    def test_page_with_no_text(self):
        mock_reader = MagicMock()
        p1 = MagicMock()
        p1.extract_text.return_value = None
        p2 = MagicMock()
        p2.extract_text.return_value = "Has text"
        mock_reader.pages = [p1, p2]

        with patch("backend.parsers.PdfReader", return_value=mock_reader):
            result = parse_pdf(b"fake pdf")

        assert result.text == "Has text"


# ── Markdown tests ──────────────────────────────────────────────────────


class TestParseMarkdown:
    def test_plain_text(self):
        content = b"Alice: Hello\nBob: Hi there"
        result = parse_markdown(content)

        assert result.format == "md"
        assert "Alice: Hello" in result.text
        assert result.detected_title is None

    def test_yaml_frontmatter(self):
        content = b'---\ntitle: "Sprint Review"\ndate: "2025-01-15"\n---\n\nTranscript content here.'
        result = parse_markdown(content)

        assert result.detected_title == "Sprint Review"
        assert result.detected_date == "2025-01-15"
        assert "Transcript content here." in result.text
        assert "---" not in result.text

    def test_structured_timestamps(self):
        content = b"[00:01:00] Alice: First point\n[00:02:00] Bob: Second point"
        result = parse_markdown(content)

        assert "[00:01:00] Alice: First point" in result.text

    def test_utf8_bom(self):
        content = b"\xef\xbb\xbfSome markdown content"
        result = parse_markdown(content)

        assert "Some markdown content" in result.text

    def test_empty_content(self):
        result = parse_markdown(b"")
        assert result.text == ""

    def test_frontmatter_without_title(self):
        content = b"---\nauthor: Anderson\n---\n\nContent here."
        result = parse_markdown(content)

        assert result.detected_title is None
        assert "Content here." in result.text


# ── Router tests ────────────────────────────────────────────────────────


class TestParseFile:
    def test_routes_vtt(self):
        vtt = b"WEBVTT\n\n00:00:01.000 --> 00:00:05.000\n<v Alice>Hello</v>\n"
        result = parse_file(vtt, "meeting.vtt")
        assert result.format == "vtt"

    def test_routes_md(self):
        result = parse_file(b"Hello world", "notes.md")
        assert result.format == "md"

    def test_unsupported_extension_raises(self):
        with pytest.raises(ValueError, match="Unsupported file type"):
            parse_file(b"content", "malware.exe")

    def test_case_insensitive_extension(self):
        result = parse_file(b"Hello world", "notes.MD")
        assert result.format == "md"

    def test_routes_pdf(self):
        mock_reader = MagicMock()
        mock_page = MagicMock()
        mock_page.extract_text.return_value = "PDF text"
        mock_reader.pages = [mock_page]

        with patch("backend.parsers.PdfReader", return_value=mock_reader):
            result = parse_file(b"fake", "report.pdf")
        assert result.format == "pdf"

    def test_routes_docx(self):
        from docx import Document

        doc = Document()
        doc.add_paragraph("Test paragraph")
        buf = io.BytesIO()
        doc.save(buf)

        result = parse_file(buf.getvalue(), "transcript.docx")
        assert result.format == "docx"
